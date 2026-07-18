from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_resume():
    doc = Document()
    
    # 设置默认字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)
    
    # 页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # ========== 姓名 ==========
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run('赵冠杰')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    name.paragraph_format.space_after = Pt(4)
    
    # ========== 求职意向 + 联系方式 ==========
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = '求职意向：AI算法/大模型应用开发实习 | 现居：山东济南 | 电话：19353177246 | 邮箱：1653094411@qq.com'
    run = contact.add_run(contact_text)
    run.font.size = Pt(9)
    contact.paragraph_format.space_after = Pt(10)
    
    # ========== 分割线 ==========
    divider = doc.add_paragraph()
    divider.add_run('_' * 85)
    divider.paragraph_format.space_after = Pt(6)
    divider.paragraph_format.space_before = Pt(0)
    
    # ========== 个人简介 ==========
    add_section_title(doc, '个人简介')
    intro = doc.add_paragraph()
    intro_text = ('软件工程专业在读，深耕大模型应用落地领域，具备完整的RAG系统与Multi-Agent架构设计与落地经验。'
                  '主导完成2个从0到1的AI系统项目，熟悉检索增强、智能体编排、效果评测全链路，具备较强的工程实现与问题优化能力。'
                  '目标在大模型应用方向持续深耕，快速为团队产出价值。')
    run = intro.add_run(intro_text)
    intro.paragraph_format.line_spacing = 1.3
    intro.paragraph_format.space_after = Pt(8)
    
    # ========== 教育背景 ==========
    add_section_title(doc, '教育背景')
    edu = doc.add_paragraph()
    run = edu.add_run('山东建筑大学')
    run.font.bold = True
    run.font.size = Pt(11)
    run = edu.add_run('        软件工程专业        本科')
    run.font.size = Pt(10.5)
    edu.add_run('        ').font.size = Pt(10.5)
    run = edu.add_run('2023.09 - 2027.06')
    run.font.size = Pt(10.5)
    edu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    edu.paragraph_format.space_after = Pt(2)
    
    course = doc.add_paragraph()
    run = course.add_run('主修课程：')
    run.font.bold = True
    course.add_run('JavaWeb应用开发、操作系统、Python程序设计、数据结构、数据库原理及应用')
    course.paragraph_format.space_after = Pt(8)
    
    # ========== 专业技能 ==========
    add_section_title(doc, '专业技能')
    
    skill1 = doc.add_paragraph(style='List Bullet')
    run = skill1.add_run('大模型应用：')
    run.font.bold = True
    skill1.add_run('精通RAG全链路产品设计与效果优化，掌握Multi-Agent架构设计、角色定义与智能路由规划，熟悉RAGAS等评测框架')
    
    skill2 = doc.add_paragraph(style='List Bullet')
    run = skill2.add_run('工程开发：')
    run.font.bold = True
    skill2.add_run('熟练使用Python、Java进行后端开发，掌握FastAPI、Streamlit等框架，具备独立完成系统从0到1搭建能力')
    
    skill3 = doc.add_paragraph(style='List Bullet')
    run = skill3.add_run('数据存储：')
    run.font.bold = True
    skill3.add_run('熟悉ChromaDB、Qdrant等向量数据库，掌握MySQL、PostgreSQL、Redis等关系型与缓存数据库')
    
    skill4 = doc.add_paragraph(style='List Bullet')
    run = skill4.add_run('语言能力：')
    run.font.bold = True
    skill4.add_run('CET-6，可熟练阅读英文技术文献，具备英文学术写作基础')
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    
    # ========== 项目经历 ==========
    add_section_title(doc, '项目经历')
    
    # 项目1
    add_project_title(doc, '智能文档检索助手（RAG问答系统）', '项目负责人', 'Python / Streamlit / LangChain / ChromaDB / MySQL')
    
    p1_desc = doc.add_paragraph()
    p1_desc.add_run('项目背景：').bold = True
    p1_desc.add_run('基于RAG技术的智能文档检索与问答系统，支持多格式文档上传、多知识库管理与流式对话，内置RAGAS评测体系实现效果量化。')
    p1_desc.paragraph_format.space_after = Pt(3)
    
    add_bullet(doc, '主导设计并实现多阶段RAG优化管线（查询改写→混合检索→LLM重排序→上下文压缩），将系统综合评测得分从0.675提升至0.896，上下文精度提升91.5%，忠实度提升28%')
    add_bullet(doc, '实现BM25关键词检索与向量语义检索的双路召回策略，结合句子级语义过滤的上下文压缩技术，显著降低无关上下文干扰')
    add_bullet(doc, '搭建完整的RAGAS自动化评测模块，覆盖上下文精确度、召回率、忠实度、答案相关性等5项核心指标，建立数据驱动的迭代机制')
    add_bullet(doc, '设计统一数据模型RawDocument中间表示层，抽象文件、网页等多源数据接入标准，新增数据源仅需开发适配器，无需修改下游处理逻辑')
    add_bullet(doc, '实现文档自动关键词提取+用户自定义标签体系，接入用户点赞/点踩反馈闭环，为持续优化提供数据支撑')
    
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    
    # 项目2
    add_project_title(doc, '全链路客户服务与工单闭环协同智能体', '核心开发', 'Python / FastAPI / LangGraph / Qdrant / PostgreSQL / Redis')
    
    p2_desc = doc.add_paragraph()
    p2_desc.add_run('项目背景：').bold = True
    p2_desc.add_run('基于Multi-Agent + LLM架构的智能客服系统，6个专业AI Agent协同工作，覆盖客户接入、意图识别、知识检索、工单创建到闭环解决全流程。')
    p2_desc.paragraph_format.space_after = Pt(3)
    
    add_bullet(doc, '设计并实现多智能体编排架构：意图分类与情绪分析并行预处理→画像增强→智能路由→回复合成的完整流水线，支持4条分支路径动态路由')
    add_bullet(doc, '搭建多路径RAG检索系统，实现查询改写+多集合并行搜索+结果融合的多路召回策略，同时检索FAQ、SOP流程、历史工单三类知识库')
    add_bullet(doc, '设计8状态+12条合法转换的工单状态机，实现工单自动分类、优先级计算与SLA时效追踪，支持P0~P3四级优先级对应1h~24h分级响应')
    add_bullet(doc, '实现LLM+规则兜底的双层情绪感知路由机制，实时监测用户情绪，识别敏感触发词时自动升级人工服务')
    add_bullet(doc, '构建知识库自进化闭环：对话结束后自动触发LLM提炼Q&A，按置信度分级（自动入库/人工审核/丢弃），支撑知识库持续迭代优化')
    add_bullet(doc, '实现检索结果PII隐私脱敏处理，保障客户数据安全合规')
    
    # 保存
    output_path = r'C:\Users\Administrator\Desktop\智能文档检索助手\赵冠杰_简历_优化版.docx'
    doc.save(output_path)
    print(f'简历已生成：{output_path}')

def add_section_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    
    # 下划线效果
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '1',
        qn('w:color'): '003366'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_project_title(doc, name, role, tech):
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.font.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p.add_run('  |  ')
    
    run = p.add_run(f'角色：{role}')
    run.font.size = Pt(9.5)
    
    p.add_run('  |  ')
    
    run = p.add_run(f'技术栈：{tech}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.25

if __name__ == '__main__':
    create_resume()
