"""文档解析器 — 支持 PDF、Word、TXT、Markdown 格式的统一解析服务"""
import mimetypes
import os
from pathlib import Path
from dataclasses import dataclass, field
from typing import List
from loguru import logger as log


# ============================================================
# 自定义异常类
# ============================================================
class DocumentProcessingError(Exception):
    """文档解析失败时抛出的异常，用于统一处理各类解析错误"""
    pass


# ============================================================
# 数据结构定义
# ============================================================
@dataclass
class PageInfo:
    """单页文档信息，用于记录文档的分页结构"""
    page_num: int           # 页码（从1开始）
    text: str               # 页面的完整文本内容
    tables: List[dict] = field(default_factory=list)  # 页面中提取的表格数据


@dataclass
class ParsedDocument:
    """文档解析后的统一输出结构，所有格式的文档最终都转换为此格式"""
    text: str               # 文档的完整文本内容（所有页合并）
    pages: List[PageInfo] = field(default_factory=list)  # 分页信息列表（保留分页结构）
    metadata: dict = field(default_factory=dict)  # 文档元数据（标题、作者、页数、文件大小等）
    images: List[dict] = field(default_factory=list)  # 已抽取的图片（本地路径、页码、格式）


# ============================================================
# 文档解析器核心类
# ============================================================
class DocumentParser:
    """文档解析器核心类，将各种格式的文档统一解析为 ParsedDocument 结构"""

    SUPPORTED_TYPES = {"pdf", "docx", "doc", "txt", "md", "markdown"}

    def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        """解析文档主入口，根据文件类型分发到对应的解析方法"""
        file_type = file_type.lower().lstrip(".")

        # 校验文件类型是否支持
        if file_type not in self.SUPPORTED_TYPES:
            raise DocumentProcessingError(
                f"不支持的文件类型: .{file_type}，支持的类型: {', '.join(self.SUPPORTED_TYPES)}"
            )

        # 校验文件是否存在
        if not os.path.exists(file_path):
            raise DocumentProcessingError(f"文件不存在: {file_path}")

        log.info(f"开始解析文档: {file_path} (类型: {file_type})")

        try:
            # 根据文件类型分发到对应的解析方法
            if file_type == "pdf":
                return self._parse_pdf(file_path)
            elif file_type in ("docx", "doc"):
                return self._parse_docx(file_path)
            elif file_type in ("txt", "md", "markdown"):
                return self._parse_text(file_path)
        except DocumentProcessingError:
            # 重新抛出自定义异常（已经是 DocumentProcessingError）
            raise
        except Exception as e:
            # 将其他异常统一包装为 DocumentProcessingError
            log.error(f"文档解析异常: {file_path} - {str(e)}")
            raise DocumentProcessingError(f"文档解析失败: {str(e)}")

    # ============================================================
# PDF 解析
# ============================================================
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """使用 PyMuPDF (fitz) 解析 PDF 文件，逐页提取文本和元数据"""
        import fitz

        doc = fitz.open(file_path)
        full_text_parts = []  # 存储每页文本
        pages = []            # 存储分页信息
        images = []
        extracted_dir = self._image_output_dir(file_path)
        seen_xrefs = set()

        # 提取 PDF 文档元数据（标题、作者、页数等）
        metadata = {
            "format": "PDF",
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "total_pages": doc.page_count,
            "file_size": os.path.getsize(file_path),
        }

        try:
            # 逐页提取文本内容
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text("text")
                full_text_parts.append(page_text)
                pages.append(PageInfo(page_num=page_num + 1, text=page_text))
                for image_index, image_info in enumerate(page.get_images(full=True), start=1):
                    xref = image_info[0]
                    if xref in seen_xrefs:
                        continue
                    seen_xrefs.add(xref)
                    extracted = doc.extract_image(xref)
                    image_bytes = extracted.get("image", b"")
                    if len(image_bytes) < 1024:
                        continue
                    extension = extracted.get("ext", "png").lower()
                    image_path = extracted_dir / f"page_{page_num + 1}_{image_index}.{extension}"
                    image_path.write_bytes(image_bytes)
                    images.append({"path": str(image_path), "page": page_num + 1, "format": extension})

            # 合并所有页的文本
            full_text = "\n\n".join(full_text_parts)
            log.info(f"PDF解析完成: {file_path} | 页数: {doc.page_count} | 总字符数: {len(full_text)}")
            return ParsedDocument(text=full_text, pages=pages, metadata=metadata, images=images)
        finally:
            # 确保文件句柄关闭
            doc.close()

    # ============================================================
# Word 文档解析
# ============================================================
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """使用 python-docx 解析 Word (.docx) 文件，提取段落和表格内容"""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = []  # 存储所有段落和表格的文本

        # 提取段落文本（跳过空段落）
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 提取表格内容（转换为文本格式，每行用 | 分隔单元格）
        for table in doc.tables:
            table_text_parts = []
            for row in table.rows:
                # 合并一行中所有单元格的文本，用 | 分隔
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text_parts.append(row_text)
            if table_text_parts:
                # 将表格的所有行合并为一个段落
                paragraphs.append("\n".join(table_text_parts))

        # 合并所有段落
        full_text = "\n\n".join(paragraphs)
        # Word 文档没有天然的分页概念，视为单页
        pages = [PageInfo(page_num=1, text=full_text)]

        # 提取 Word 文档元数据
        metadata = {
            "format": "DOCX",
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "paragraph_count": len(doc.paragraphs),
            "file_size": os.path.getsize(file_path),
        }

        images = []
        extracted_dir = self._image_output_dir(file_path)
        seen_parts = set()
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype or rel.is_external:
                continue
            image_part = rel.target_part
            if image_part.partname in seen_parts:
                continue
            seen_parts.add(image_part.partname)
            content_type = image_part.content_type
            extension = mimetypes.guess_extension(content_type) or ".png"
            image_path = extracted_dir / f"image_{len(images) + 1}{extension}"
            image_path.write_bytes(image_part.blob)
            images.append({"path": str(image_path), "page": 1, "format": content_type})

        log.info(f"DOCX解析完成: {file_path} | 段落数: {len(doc.paragraphs)} | 总字符数: {len(full_text)}")
        return ParsedDocument(text=full_text, pages=pages, metadata=metadata, images=images)

    @staticmethod
    def _image_output_dir(file_path: str) -> Path:
        """为源文件的抽取图片创建稳定、隔离的本地目录。"""
        directory = Path(file_path).with_suffix("").parent / "extracted_images" / Path(file_path).stem
        directory.mkdir(parents=True, exist_ok=True)
        return directory

    # ============================================================
# 纯文本/Markdown 解析
# ============================================================
    def _parse_text(self, file_path: str) -> ParsedDocument:
        """解析纯文本或 Markdown 文件，按行分段模拟分页"""
        # 以 UTF-8 编码读取文件，遇到编码错误时替换为替换字符
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        file_type = os.path.splitext(file_path)[1].lower().replace(".", "")
        lines = content.split("\n")

        # 将长文本按行分段（每200行作为一个"页面"，便于后续处理）
        pages = []
        chunk_size = 200
        for i in range(0, len(lines), chunk_size):
            chunk_lines = lines[i:i + chunk_size]
            pages.append(PageInfo(page_num=len(pages) + 1, text="\n".join(chunk_lines)))

        # 提取文本文件元数据
        metadata = {
            "format": file_type.upper(),
            "line_count": len(lines),
            "file_size": os.path.getsize(file_path),
        }

        log.info(f"{file_type.upper()}解析完成: {file_path} | 行数: {len(lines)} | 总字符数: {len(content)}")
        return ParsedDocument(text=content, pages=pages, metadata=metadata)


# ============================================================
# 单例实例
# ============================================================
# 创建全局单例，其他模块直接导入使用，避免重复创建解析器实例
document_parser = DocumentParser()
